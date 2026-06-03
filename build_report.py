#!/usr/bin/env python3
"""Build an easy-to-read Word report from results.json + figures."""
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

R = json.load(open("results.json"))


def pfmt(p):
    return "< .001" if p < 0.001 else f"= {p:.3f}"


def es_word(r):
    a = abs(r)
    return "negligible" if a < 0.1 else "small" if a < 0.3 else "moderate" if a < 0.5 else "large"


doc = Document()
# base style
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

t = doc.add_heading("Does conversation modality (VR vs Text) affect self-disclosure?", level=0)
sub = doc.add_paragraph("AI Companion Lab Study — disclosure-depth analysis")
sub.runs[0].italic = True
doc.add_paragraph("Prepared 3 June 2026").runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ---------- TL;DR ----------
doc.add_heading("The short version", level=1)
for b in [
    "Yes — there is a difference: participants disclosed MORE deeply in the Text condition than in VR.",
    "It looked like “no difference” at first because VR participants talked far more (≈322 vs ≈149 words). "
    "Sheer volume nudges depth scores up, which masked the text advantage.",
    "Once we account for word count (and topic sensitivity), the text advantage is large and highly significant: "
    f"odds of deeper disclosure are about 5× higher on the 0–3 scale and 7.7× higher on the 0–6 scale (both p < .001).",
    "Plain-English mechanism: typing is more deliberate and lower-pressure, so people open up more per word; "
    "speaking produces lots of words but much of it is filler and external description, not personal disclosure.",
]:
    doc.add_paragraph(b, style="List Bullet")

# ---------- what we measured ----------
doc.add_heading("What we measured", level=1)
doc.add_paragraph(
    "Every participant's full conversation transcript (n = 187 with codeable data) was scored for disclosure "
    "depth on two rubrics: Rubric 1 (0–3: none / surface / moderate / deep) and Rubric 2 (0–6, a finer "
    "split of the same construct). Scoring followed the project rubric and the multi-turn rules — content "
    "only (never length), crediting only what the participant originated (not the AI's emotional reframes). "
    "We compared the two randomised conditions: VR (spoken) vs Text (chat)."
)

# ---------- headline numbers ----------
doc.add_heading("Headline numbers", level=1)
doc.add_paragraph(f"Group sizes: VR n = {R['n_vr']}, Text n = {R['n_text']}.")
tbl = doc.add_table(rows=1, cols=6)
tbl.style = "Light Grid Accent 1"
for i, h in enumerate(["Measure", "VR (mean)", "Text (mean)", "Test", "p", "Effect size"]):
    tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
rows = [
    ("Disclosure depth R1 (0–3)", R["R1"], "Mann–Whitney"),
    ("Disclosure depth R2 (0–6)", R["R2"], "Mann–Whitney"),
    ("Word count", R["WC"], "Mann–Whitney"),
]
for name, s, test in rows:
    c = tbl.add_row().cells
    c[0].text = name
    c[1].text = f"{s['vr_mean']:.2f}"
    c[2].text = f"{s['tx_mean']:.2f}"
    c[3].text = test
    c[4].text = pfmt(s["p"])
    c[5].text = f"r = {s['r']:+.2f} ({es_word(s['r'])})"
doc.add_paragraph(
    "Read raw, the depth difference is small — significant on the finer R2 scale (p = .025) and just short of "
    "significance on R1 (p = .076). The word-count gap, by contrast, is enormous (p < .001): VR participants "
    "produced more than twice as many words."
).runs[0].italic = True

# ---------- the key insight ----------
doc.add_heading("Why it looked like “no difference” (the important part)", level=1)
doc.add_paragraph(
    "Word count and disclosure depth are positively related — a longer transcript has more opportunity for a "
    "personal signal to appear. Because VR transcripts are much longer, VR scores got an artificial “volume "
    "boost” that cancelled out the genuine text advantage. This is a classic suppression effect: the real "
    "effect is hidden until you control for the nuisance variable."
)
doc.add_paragraph(
    "We ran an ordinal logistic regression predicting disclosure depth from modality, while controlling for word "
    "count and topic sensitivity. Odds ratios (OR) above 1 mean greater odds of a deeper disclosure level."
)


def reg_table(title, ord_res):
    doc.add_paragraph().add_run(title).bold = True
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["Predictor", "Odds ratio", "95% CI", "p"]):
        t.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    labels = {"text": "Text vs VR", "ts": "Topic sensitivity (per level)", "wc_z": "Word count (per SD)"}
    for term in ["text", "ts", "wc_z"]:
        v = ord_res[term]
        c = t.add_row().cells
        c[0].text = labels[term]
        c[1].text = f"{v['OR']:.2f}"
        c[2].text = f"{v['ci'][0]:.2f} – {v['ci'][1]:.2f}"
        c[3].text = pfmt(v["p"])


reg_table("Rubric 1 (0–3):", R["ord_R1"])
reg_table("Rubric 2 (0–6):", R["ord_R2"])
doc.add_paragraph(
    f"Adjusted for word count and topic sensitivity, Text raises the odds of deeper disclosure about "
    f"{R['ord_R1']['text']['OR']:.0f}× on R1 and {R['ord_R2']['text']['OR']:.0f}× on R2 (both p < .001). "
    "Topic sensitivity and word count are also significant positive predictors, as expected."
)

# ---------- figures ----------
doc.add_heading("Figures", level=1)
for fn, cap in [
    ("fig1_depth_by_modality.png", "Figure 1. Disclosure-depth distribution (Rubric 1) by modality. Text shifts toward the Moderate level; VR has more None/Surface."),
    ("fig2_R2_box.png", "Figure 2. Finer-grained depth (Rubric 2) by modality. Text's median is higher."),
    ("fig3_wordcount_vs_depth.png", "Figure 3. Word count vs disclosure depth. VR transcripts (blue) are far longer but not deeper — length is not depth."),
    ("fig4_wordcount.png", "Figure 4. Mean word count by modality (± SE). VR produces more than twice as many words."),
]:
    doc.add_picture(fn, width=Inches(5.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = doc.add_paragraph(cap)
    cp.runs[0].italic = True
    cp.runs[0].font.size = Pt(9.5)

# ---------- interpretation ----------
doc.add_heading("What it means", level=1)
doc.add_paragraph(
    "Your instinct was right: people disclose more in text. The effect was hidden in a simple comparison because "
    "the voice condition generates a high volume of words — filler, false starts, and external narration — "
    "that inflates raw scores without reflecting genuine personal disclosure. Per unit of content, text "
    "participants reveal substantially more. This fits the wider computer-mediated-communication literature: "
    "text lowers social presence and lets people edit and pace themselves, which encourages self-disclosure, "
    "whereas real-time voice raises self-presentation pressure."
)

# ---------- caveats ----------
doc.add_heading("Caveats & recommended next steps", level=1)
for b in [
    "Depth scores are AI-coded by a single rater. Before publishing, have a human second-code a random subset "
    "(e.g. 20–30 transcripts) and report inter-rater reliability (weighted kappa). The justification column "
    "in the spreadsheet supports this.",
    "Raw between-group effect sizes are small (r ≈ .14–.18); the strong effect emerges only after adjusting "
    "for word count. Report both the unadjusted and adjusted results honestly.",
    "Consider pre-registering / clearly labelling the word-count-adjusted model as the primary test, since it is "
    "the theoretically meaningful comparison (depth independent of verbosity).",
    "Optional robustness checks: rerun excluding the No-Disclosure (0) cases, and excluding off-task transcripts "
    "flagged in the coding notes (e.g. 15253, 14854).",
]:
    doc.add_paragraph(b, style="List Bullet")

# ---------- data location ----------
doc.add_heading("Where the data lives", level=1)
doc.add_paragraph(
    "The SPSS file now contains four new variables: DD_R1 (0–3), DD_R2 (0–6), WordCount, and "
    "TopicSensitivity (0 None – 3 High), plus the full Transcript text. Condition is labelled 1 = VR, "
    "2 = Text. Per-participant scores with justifications are in Transcripts_participant_only.xlsx."
)

doc.save("Modality_Disclosure_Report.docx")
print("wrote Modality_Disclosure_Report.docx")
