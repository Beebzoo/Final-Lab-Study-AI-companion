#!/usr/bin/env python3
"""Mediation diagnostics + recommendations report for Study 5 (filter on)."""
import numpy as np, pandas as pd, pyreadstat, statsmodels.api as sm
from scipy import stats
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = "/root/.claude/uploads/df0dbb36-4d78-424d-bd29-c1d29cdf89f4/819687d3-AI_Companions__Study_5__no_attention_checks.sav"
df, _ = pyreadstat.read_sav(SRC)
d = df[df["filter_$"] == 1].copy()
d["text"] = (d["Condition"] == 2).astype(float)
d["SPE"] = d[["SPE_1", "SPE_2", "SPE_3"]].mean(axis=1)
ITEMS = {"IC": ["IC_1", "IC_2", "IC_3"], "AF": ["AF_1", "AF_2", "AF_3"], "PP": ["PP_1", "PP_2", "PP_3", "PP_4"]}


def alpha(items):
    X = d[items].dropna(); k = X.shape[1]
    return k / (k - 1) * (1 - X.var(ddof=1).sum() / X.sum(axis=1).var(ddof=1))


def apath(M):
    m = sm.OLS(d[M], sm.add_constant(d["text"]), missing="drop").fit()
    return d[d.text == 0][M].mean(), d[d.text == 1][M].mean(), m.pvalues["text"]


def bpath(M, DV):
    m = sm.OLS(d[DV], sm.add_constant(d[["text", M]]), missing="drop").fit()
    return m.params[M], m.pvalues[M]


def p(x):
    return "< .001" if x < 0.001 else f"= {x:.3f}"


mc_ok = (d["MC"] == d["Condition"]).mean()
# word-count suppression
aW = sm.OLS(d["WC"], sm.add_constant(d["text"]), missing="drop").fit()
bW = sm.OLS(d["DD_L"], sm.add_constant(d[["text", "WC"]]), missing="drop").fit()

doc = Document()
doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(11)
doc.add_heading("Why the mediators (IC, AF, PP) don't work — and how to fix it", level=0)
s = doc.add_paragraph("AI Companion Study 5 — mediation diagnostics (analysis filter applied, n = %d)" % len(d.dropna(subset=["DD_L"])))
s.runs[0].italic = True
doc.add_paragraph("Prepared 3 June 2026").runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_heading("Bottom line", level=1)
for b in [
    "The single decisive reason: the modality manipulation does NOT change any of the proposed mediators "
    "(the a-path is null for IC, AF, and PP). A mediator can only transmit an effect if the conditions differ "
    "on it — these don't, so no indirect effect is possible.",
    "This is not a failed manipulation: 100%% of participants correctly recalled their format. People knew they "
    "were in VR vs Text; it simply did not change their felt control, flow, or presence.",
    "Therefore this should be reported as an informative null (the disclosure effect is NOT explained by these "
    "experiences), not 'fixed' by searching for a significant mediator.",
]:
    doc.add_paragraph(b.replace("100%%", "100%"), style="List Bullet")

# a-path table
doc.add_heading("Diagnostic 1 — a-paths (does modality move the mediator?)", level=1)
t = doc.add_table(rows=1, cols=5); t.style = "Light Grid Accent 1"
for i, h in enumerate(["Mediator", "VR mean", "Text mean", "a-path p", "Verdict"]):
    t.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
for M, name in [("IC", "IC – information control"), ("AF", "AF – flow/absorption"),
                ("PP", "PP – presence/absorption"), ("SPE", "SPE – self-presentation effort")]:
    vr, tx, pv = apath(M); c = t.add_row().cells
    c[0].text = name; c[1].text = f"{vr:.2f}"; c[2].text = f"{tx:.2f}"; c[3].text = p(pv)
    c[4].text = "no difference" if pv > .05 else "differs"
doc.add_paragraph("All a-paths are non-significant: modality does not shift any mediator. This is the binding constraint.").runs[0].italic = True

# b-path table
doc.add_heading("Diagnostic 2 — b-paths (does the mediator predict disclosure?)", level=1)
doc.add_paragraph("Two possible outcomes: DD_L = AI-coded behavioural depth (0–6); DC = self-reported disclosure.")
t = doc.add_table(rows=1, cols=5); t.style = "Light Grid Accent 1"
for i, h in enumerate(["Mediator", "→ DD_L (coded)", "p", "→ DC (self-report)", "p"]):
    t.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
for M in ["IC", "AF", "PP"]:
    bd, pd_ = bpath(M, "DD_L"); bc, pc = bpath(M, "DC"); c = t.add_row().cells
    c[0].text = M; c[1].text = f"b = {bd:+.2f}"; c[2].text = p(pd_); c[3].text = f"b = {bc:+.2f}"; c[4].text = p(pc)
doc.add_paragraph(
    "AF and PP strongly predict self-reported disclosure (DC) but not coded depth (DD_L); IC predicts neither. "
    "So AF/PP are genuine correlates of felt disclosure — just not mediators of the modality effect, because "
    "modality doesn't move them."
).runs[0].italic = True

# reliability + redundancy
doc.add_heading("Diagnostic 3 — measurement quality", level=1)
for b in [
    f"Reliability (Cronbach's alpha): IC = {alpha(ITEMS['IC']):.2f} (good), PP = {alpha(ITEMS['PP']):.2f} (good), "
    f"AF = {alpha(ITEMS['AF']):.2f} (POOR, below .70). AF is an unreliable scale, which attenuates its relationships.",
    f"AF and PP correlate r = {d[['AF','PP']].corr().iloc[0,1]:.2f} — they measure nearly the same construct, so "
    "they should not be entered as two separate mediators.",
    "IC sits near the ceiling (~5.8 / 7) in both conditions, leaving little variance to predict anything.",
    f"Coded depth (DD_L) and self-report (DC) correlate only r = {d[['DD_L','DC']].corr().iloc[0,1]:.2f} — they are "
    "different constructs and must not be treated interchangeably.",
]:
    doc.add_paragraph(b, style="List Bullet")

# word count suppression
doc.add_heading("Diagnostic 4 — the one pathway that is active: word count (suppression)", level=1)
doc.add_paragraph(
    f"Modality strongly moves word count (Text → {aW.params['text']:+.0f} words, p < .001), and word count predicts "
    f"depth (b = {bW.params['WC']:+.4f}, p < .001). But the indirect effect ({aW.params['text']*bW.params['WC']:+.2f}) "
    f"is opposite in sign to the direct effect ({bW.params['text']:+.2f}). This is suppression / inconsistent "
    "mediation, not classic mediation: text participants write fewer words yet disclose more per word."
)

# recommendations
doc.add_heading("How to fix each problem", level=1)
recs = [
    ("Null a-paths (the main issue)", "Report as an informative null — the disclosure effect is not carried by "
     "control/flow/presence/effort. Do NOT keep swapping mediators until one is significant (that capitalises on "
     "chance). Pre-register the revised model."),
    ("AF reliability (alpha = .57)", "Don't report AF as-is. Merge AF and PP into one 'absorption/presence' factor "
     "(they correlate .76); confirm with EFA/CFA and report the combined alpha. In future, use a validated flow/"
     "presence scale."),
    ("AF/PP redundancy", "Enter a single immersion composite, not two near-identical mediators (avoids collinearity)."),
    ("IC ceiling / no effect", "Treat IC as a non-starter here (no variance, no a- or b-path) and note the ceiling."),
    ("DV ambiguity (DD vs DC)", "Choose one primary outcome a priori from theory. AF/PP belong to the subjective (DC) "
     "story; word count belongs to the behavioural (DD) story. Keep them separate."),
    ("Noisy coded depth", "Add a human second-coder for DD on a subset and report inter-rater reliability; "
     "single-rater AI coding attenuates every b-path to DD."),
    ("Word-count pathway", "Model it as suppression with bootstrapped indirect effects (e.g. PROCESS / semopy), and "
     "describe it correctly: fewer words but deeper disclosure in text."),
    ("If experiential mechanism is central (new data)", "Pilot the VR manipulation to confirm it actually raises "
     "presence (here it didn't), and measure the likely real driver of text disclosure: editability/revisability, "
     "asynchronous deliberation, reduced social presence, response latency."),
]
for title, body in recs:
    pgh = doc.add_paragraph(style="List Bullet")
    pgh.add_run(title + ": ").bold = True
    pgh.add_run(body)

doc.add_heading("Recommended reframe", level=1)
for b in [
    "Behavioural story: Modality -> word count (suppressor) -> coded depth; text raises depth directly, independent of verbosity.",
    "Subjective story: presence/absorption -> felt disclosure (DC) as an individual-difference predictor, not a mediator of modality.",
    "Presence/flow as an OUTCOME: 'Did VR increase immersion over text?' Here, no — itself worth reporting.",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.save("Study5_Mediation_Diagnostics.docx")
print("wrote Study5_Mediation_Diagnostics.docx")
