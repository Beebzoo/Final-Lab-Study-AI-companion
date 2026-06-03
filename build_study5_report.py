#!/usr/bin/env python3
"""Build an easy-to-read Word report for Study 5 (filter on)."""
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

R = json.load(open("study5_results.json"))


def pfmt(p):
    return "< .001" if p < 0.001 else f"= {p:.3f}"


def es_word(r):
    a = abs(r)
    return "negligible" if a < 0.1 else "small" if a < 0.3 else "moderate" if a < 0.5 else "large"


doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

doc.add_heading("Study 5 — Does modality (VR vs Text) affect self-disclosure?", level=0)
p = doc.add_paragraph("AI Companion Study 5 — disclosure-depth analysis (analysis filter applied)")
p.runs[0].italic = True
doc.add_paragraph("Prepared 3 June 2026").runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_heading("The short version", level=1)
for b in [
    f"Filter applied: {R['n_excluded_by_filter']} case(s) excluded by filter_$, leaving {R['n_after_filter']}; "
    f"{R['n_analysed']} had complete disclosure scores and were analysed ({R['n_vr']} VR, {R['n_text']} Text).",
    "Same result as the Lab Study: participants disclosed MORE deeply in Text than in VR.",
    "It again looks like “no difference” at first glance because VR participants talked far more "
    f"(≈{R['WC']['vr_mean']:.0f} vs ≈{R['WC']['tx_mean']:.0f} words), which inflates raw depth scores and masks the text advantage.",
    "Adjusted for word count, the Text advantage is large and highly significant: odds of deeper disclosure are "
    f"about {R['ord_R1']['text']['OR']:.0f}× higher on the 0–3 scale and {R['ord_R2']['text']['OR']:.0f}× higher on the 0–6 scale (both p < .001).",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("What we measured", level=1)
doc.add_paragraph(
    "Disclosure depth was taken from the scores already in the dataset: DD_S (0–3) and DD_L (0–6, a finer split "
    "of the same construct), with WC = participant word count. We compared the two randomised conditions, "
    "VR (spoken) vs Text (chat), on the filtered analysis sample."
)

doc.add_heading("Headline numbers", level=1)
doc.add_paragraph(f"Group sizes (filter on): VR n = {R['n_vr']}, Text n = {R['n_text']}.")
tbl = doc.add_table(rows=1, cols=6); tbl.style = "Light Grid Accent 1"
for i, h in enumerate(["Measure", "VR (mean)", "Text (mean)", "Test", "p", "Effect size"]):
    tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
for name, s in [("Disclosure depth DD_S (0–3)", R["R1"]),
                ("Disclosure depth DD_L (0–6)", R["R2"]),
                ("Word count", R["WC"])]:
    c = tbl.add_row().cells
    c[0].text = name; c[1].text = f"{s['vr_mean']:.2f}"; c[2].text = f"{s['tx_mean']:.2f}"
    c[3].text = "Mann–Whitney"; c[4].text = pfmt(s["p"]); c[5].text = f"r = {s['r']:+.2f} ({es_word(s['r'])})"
doc.add_paragraph(
    "Raw, the depth difference is small — significant on DD_L (p = .030) and just short on DD_S (p = .093) — "
    "while the word-count gap is enormous (p < .001)."
).runs[0].italic = True

doc.add_heading("Why it looks like “no difference” (the important part)", level=1)
doc.add_paragraph(
    "Longer transcripts have more opportunity for a personal signal to appear, so word count is positively related "
    "to depth. Because VR transcripts are far longer, VR scores get a “volume boost” that cancels out the genuine "
    "text advantage — a classic suppression effect. Controlling for word count reveals the real effect."
)
doc.add_paragraph("Ordinal logistic regression predicting disclosure depth from modality, controlling for word count. "
                  "Odds ratios (OR) above 1 mean greater odds of a deeper disclosure level.")


def reg_table(title, ord_res):
    doc.add_paragraph().add_run(title).bold = True
    t = doc.add_table(rows=1, cols=4); t.style = "Light Grid Accent 1"
    for i, h in enumerate(["Predictor", "Odds ratio", "95% CI", "p"]):
        t.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    labels = {"text": "Text vs VR", "wc_z": "Word count (per SD)"}
    for term in ["text", "wc_z"]:
        v = ord_res[term]; c = t.add_row().cells
        c[0].text = labels[term]; c[1].text = f"{v['OR']:.2f}"
        c[2].text = f"{v['ci'][0]:.2f} – {v['ci'][1]:.2f}"; c[3].text = pfmt(v["p"])


reg_table("DD_S (0–3):", R["ord_R1"])
reg_table("DD_L (0–6):", R["ord_R2"])
doc.add_paragraph(
    f"Adjusted for word count, Text raises the odds of deeper disclosure about {R['ord_R1']['text']['OR']:.0f}× on "
    f"DD_S and {R['ord_R2']['text']['OR']:.0f}× on DD_L (both p < .001). Word count is also a significant positive predictor."
)

doc.add_heading("Figures", level=1)
for fn, cap in [
    ("s5_fig1_depth_by_modality.png", "Figure 1. Disclosure-depth distribution (DD_S) by modality."),
    ("s5_fig2_R2_box.png", "Figure 2. Finer-grained depth (DD_L) by modality; Text's median is higher."),
    ("s5_fig3_wordcount_vs_depth.png", "Figure 3. Word count vs depth: VR is longer (blue) but not deeper."),
    ("s5_fig4_wordcount.png", "Figure 4. Mean word count by modality (± SE)."),
]:
    doc.add_picture(fn, width=Inches(5.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = doc.add_paragraph(cap); cp.runs[0].italic = True; cp.runs[0].font.size = Pt(9.5)

doc.add_heading("What it means", level=1)
doc.add_paragraph(
    "Study 5 replicates the Lab Study: people disclose more in text. The effect hides in a simple comparison "
    "because the voice condition generates a high volume of words (filler, false starts, external narration) that "
    "inflates raw scores without reflecting genuine personal disclosure. Per unit of content, text participants "
    "reveal substantially more — consistent with the idea that text lowers social-presence pressure and lets people "
    "edit and pace their self-disclosure."
)

doc.add_heading("Caveats & notes", level=1)
for b in [
    "The analysis filter (filter_$) was applied: 2 cases excluded; 177 analysed.",
    "This file had no topic-sensitivity variable, so the model controls for word count only (the key suppressor). "
    "I can add a topic-sensitivity control if you code/supply that variable.",
    "Depth scores (DD_S/DD_L) were taken as given in the file; if they are AI-coded, a human reliability check on a "
    "subset is still recommended before publication.",
    "Raw between-group effects are small (r ≈ .13–.18); the strong effect is the word-count-adjusted model. "
    "Report both honestly.",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.save("Study5_Disclosure_Report.docx")
print("wrote Study5_Disclosure_Report.docx")
